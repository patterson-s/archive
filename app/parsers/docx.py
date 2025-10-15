from docx import Document
from typing import Dict

def parse_docx(file_path: str) -> Dict[str, str]:
    doc = Document(file_path)
    
    text_parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            if para.style.name.startswith('Heading'):
                text_parts.append(f"\n{text}\n")
            else:
                text_parts.append(text)
    
    full_text = "\n".join(text_parts)
    
    metadata = {
        "num_paragraphs": len([p for p in doc.paragraphs if p.text.strip()]),
        "content": full_text
    }
    
    return metadata